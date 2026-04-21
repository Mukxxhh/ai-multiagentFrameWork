"""
File Parser Utility - Handles parsing of various file formats.
"""
import os
import json
from io import StringIO
import pandas as pd
import pdfplumber
from docx import Document
from PyPDF2 import PdfReader
import chardet
from typing import Dict, Any, List, Tuple
import traceback

try:
    import win32com.client  # type: ignore[import-not-found]
except ImportError:
    win32com = None


class FileParser:
    """Parse various file formats and extract content."""

    SUPPORTED_FORMATS = ['csv', 'xlsx', 'xls', 'json', 'pdf', 'txt', 'doc', 'docx', 'xml', 'parquet']

    def __init__(self):
        self.data = None
        self.text_content = ""
        self.metadata = {}
        self.file_type = None

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file and return structured data.

        Args:
            file_path: Path to the file to parse

        Returns:
            Dictionary containing parsed data, text content, and metadata
        """
        self.file_type = self._get_file_extension(file_path)

        if self.file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {self.file_type}")

        self.metadata = {
            'file_name': os.path.basename(file_path),
            'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
            'file_type': self.file_type
        }

        # Parse based on file type
        parser_map = {
            'csv': self._parse_csv,
            'xlsx': self._parse_excel,
            'xls': self._parse_excel,
            'json': self._parse_json,
            'pdf': self._parse_pdf,
            'txt': self._parse_text,
            'doc': self._parse_doc,
            'docx': self._parse_docx,
            'xml': self._parse_xml,
            'parquet': self._parse_parquet
        }

        try:
            parser_map[self.file_type](file_path)
        except Exception as e:
            raise Exception(f"Error parsing {self.file_type} file: {str(e)}\n{traceback.format_exc()}")

        return {
            'data': self.data,
            'text_content': self.text_content,
            'metadata': self.metadata,
            'file_type': self.file_type,
            'has_numerical_data': self._has_numerical_data()
        }

    def parse_text_input(self, raw_text: str, source_name: str = "direct_input.txt") -> Dict[str, Any]:
        """
        Parse direct user text input and detect whether it behaves like
        tabular/numerical data or general text.
        """
        self.data = None
        self.text_content = raw_text or ""
        self.file_type = "text_input"
        self.metadata = {
            "file_name": source_name,
            "file_size_mb": len((raw_text or "").encode("utf-8")) / (1024 * 1024),
            "file_type": self.file_type,
            "line_count": len((raw_text or "").splitlines()),
        }

        stripped = (raw_text or "").strip()

        if not stripped:
            return {
                "data": None,
                "text_content": "",
                "metadata": self.metadata,
                "file_type": self.file_type,
                "has_numerical_data": False,
            }

        # Try JSON first.
        try:
            json_data = json.loads(stripped)
            if isinstance(json_data, list) and all(isinstance(item, dict) for item in json_data):
                self.data = pd.DataFrame(json_data)
                self.text_content = self.data.to_string()
                self.metadata.update({
                    "rows": len(self.data),
                    "columns": list(self.data.columns),
                    "detected_format": "json_array",
                })
            elif isinstance(json_data, dict):
                normalized = pd.json_normalize(json_data)
                if not normalized.empty and len(normalized.columns) > 1:
                    self.data = normalized
                    self.text_content = normalized.to_string()
                    self.metadata.update({
                        "rows": len(normalized),
                        "columns": list(normalized.columns),
                        "detected_format": "json_object",
                    })
        except Exception:
            pass

        # Try CSV/TSV style tabular text.
        if self.data is None:
            for separator, label in [(",", "csv"), ("\t", "tsv"), (";", "semicolon_delimited")]:
                try:
                    frame = pd.read_csv(StringIO(stripped), sep=separator)
                    if not frame.empty and len(frame.columns) > 1:
                        self.data = frame
                        self.text_content = frame.to_string()
                        self.metadata.update({
                            "rows": len(frame),
                            "columns": list(frame.columns),
                            "detected_format": label,
                        })
                        break
                except Exception:
                    continue

        if self.data is None:
            self.metadata["detected_format"] = "plain_text"

        return {
            "data": self.data,
            "text_content": self.text_content,
            "metadata": self.metadata,
            "file_type": self.file_type,
            "has_numerical_data": self._has_numerical_data(),
        }

    def _get_file_extension(self, file_path: str) -> str:
        """Extract file extension."""
        return os.path.splitext(file_path)[1].lower().lstrip('.')

    def _parse_csv(self, file_path: str):
        """Parse CSV files with encoding detection."""
        # Detect encoding
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
            encoding = result['encoding']

        self.data = pd.read_csv(file_path, encoding=encoding)
        self.text_content = self.data.to_string()
        self.metadata.update({
            'rows': len(self.data),
            'columns': list(self.data.columns),
            'encoding': encoding
        })

    def _parse_excel(self, file_path: str):
        """Parse Excel files."""
        self.data = pd.read_excel(file_path)
        self.text_content = self.data.to_string()
        self.metadata.update({
            'rows': len(self.data),
            'columns': list(self.data.columns),
            'sheets': pd.ExcelFile(file_path).sheet_names
        })

    def _parse_json(self, file_path: str):
        """Parse JSON files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)

        # Try to convert to DataFrame if it's tabular
        if isinstance(json_data, list) and all(isinstance(item, dict) for item in json_data):
            self.data = pd.DataFrame(json_data)
            self.text_content = self.data.to_string()
            self.metadata.update({
                'rows': len(self.data),
                'columns': list(self.data.columns),
                'json_type': 'array_of_objects'
            })
        else:
            self.data = None
            self.text_content = json.dumps(json_data, indent=2)
            self.metadata.update({
                'json_type': type(json_data).__name__
            })

    def _parse_pdf(self, file_path: str):
        """Parse PDF files."""
        text_parts = []
        tables = []

        with pdfplumber.open(file_path) as pdf:
            self.metadata.update({
                'pages': len(pdf.pages)
            })

            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        try:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            tables.append(df)
                        except:
                            pass

        self.text_content = "\n\n".join(text_parts)

        # If tables found, use the first one as primary data
        if tables:
            self.data = tables[0]
            self.metadata['tables_found'] = len(tables)

    def _parse_text(self, file_path: str):
        """Parse plain text files."""
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
            encoding = result['encoding']

        with open(file_path, 'r', encoding=encoding) as f:
            self.text_content = f.read()

        self.metadata['encoding'] = encoding
        self.metadata['line_count'] = len(self.text_content.split('\n'))

    def _parse_docx(self, file_path: str):
        """Parse Word documents."""
        doc = Document(file_path)

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        self.text_content = "\n\n".join(paragraphs)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            if table_data:
                df = pd.DataFrame(table_data[1:], columns=table_data[0])
                tables.append(df)

        if tables:
            self.data = tables[0]

        self.metadata.update({
            'paragraphs': len(paragraphs),
            'tables': len(tables)
        })

    def _parse_doc(self, file_path: str):
        """Parse legacy Word .doc files through Microsoft Word on Windows."""
        if win32com is None:
            raise RuntimeError("Legacy .doc support requires pywin32 / win32com on Windows.")

        word_app = None
        document = None
        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            document = word_app.Documents.Open(os.path.abspath(file_path), ReadOnly=True)

            paragraphs = []
            for paragraph in document.Paragraphs:
                text = str(paragraph.Range.Text).strip()
                if text:
                    paragraphs.append(text)

            self.text_content = "\n\n".join(paragraphs).strip()
            self.metadata.update({
                'paragraphs': len(paragraphs),
                'tables': int(document.Tables.Count),
                'parsed_with': 'microsoft_word_com',
            })
        except Exception as exc:
            raise RuntimeError(
                "Could not read the .doc file. Please make sure Microsoft Word is installed on this Windows machine."
            ) from exc
        finally:
            if document is not None:
                document.Close(False)
            if word_app is not None:
                word_app.Quit()

    def _parse_xml(self, file_path: str):
        """Parse XML files."""
        self.data = pd.read_xml(file_path)
        self.text_content = self.data.to_string()
        self.metadata.update({
            'rows': len(self.data),
            'columns': list(self.data.columns)
        })

    def _parse_parquet(self, file_path: str):
        """Parse Parquet files."""
        self.data = pd.read_parquet(file_path)
        self.text_content = self.data.to_string()
        self.metadata.update({
            'rows': len(self.data),
            'columns': list(self.data.columns)
        })

    def _has_numerical_data(self) -> bool:
        """Check if the data contains numerical columns."""
        if self.data is None:
            return False

        if isinstance(self.data, pd.DataFrame):
            numeric_cols = self.data.select_dtypes(include=['number']).columns
            return len(numeric_cols) > 0

        return False

    def get_numerical_columns(self) -> List[str]:
        """Get list of numerical column names."""
        if self.data is None or not isinstance(self.data, pd.DataFrame):
            return []

        return list(self.data.select_dtypes(include=['number']).columns)

    def get_categorical_columns(self) -> List[str]:
        """Get list of categorical column names."""
        if self.data is None or not isinstance(self.data, pd.DataFrame):
            return []

        return list(self.data.select_dtypes(include=['object', 'category']).columns)
